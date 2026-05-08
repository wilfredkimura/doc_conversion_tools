import os
from pathlib import Path
import markdown
from docx import Document
from bs4 import BeautifulSoup
from .base import BaseConverter

class MarkdownToDocxConverter(BaseConverter):
    """
    Converts Markdown (.md) files to Word (.docx).
    Uses a basic HTML intermediate step for parsing.
    """

    def __init__(self):
        pass

    def convert(self, input_path: Path, output_dir: Path) -> Path:
        if not input_path.exists():
            raise FileNotFoundError(f"File not found: {input_path}")

        base_name = input_path.stem
        output_docx = output_dir / f"{base_name}.docx"

        try:
            with open(input_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            # Convert MD to HTML
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')

            # Create Word Doc
            doc = Document()
            
            for element in soup.find_all(True):
                if element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name == 'h1':
                    doc.add_heading(element.get_text(), level=0)
                elif element.name == 'h2':
                    doc.add_heading(element.get_text(), level=1)
                elif element.name == 'h3':
                    doc.add_heading(element.get_text(), level=2)
                elif element.name == 'h4':
                    doc.add_heading(element.get_text(), level=3)
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text(), style='List Bullet')
                elif element.name == 'ol':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text(), style='List Number')
                elif element.name == 'pre':
                    doc.add_paragraph(element.get_text(), style='Quote')
                
                # Note: This is a basic implementation. 
                # Complex nested elements might need more recursion.

            doc.save(str(output_docx))
            return output_docx
            
        except Exception as e:
            raise RuntimeError(f"Failed to convert {input_path.name}: {str(e)}")
